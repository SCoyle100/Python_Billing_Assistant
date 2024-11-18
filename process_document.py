import sys
from PyQt5.QtWidgets import QApplication, QMessageBox
import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import os
from dotenv import load_dotenv
import invoice
import fitz  # PyMuPDF
from pdf_to_docx import PDFConverter
import logging
import win32com.client as win32

logging.basicConfig(level=logging.DEBUG)

# Function to get the path of the docx file
def get_docx_path(pdf_filename):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(script_dir, "output")

    if not os.path.exists(output_dir):
        logging.error(f"Output directory does not exist: {output_dir}")
        return None

    base_name = os.path.basename(pdf_filename)
    file_name, _ = os.path.splitext(base_name)
    docx_path = os.path.join(output_dir, f"{file_name}.docx")

    if not os.path.exists(docx_path):
        logging.error(f"DOCX file does not exist: {docx_path}")
        return None

    logging.debug(f"Retrieved docx path: {docx_path}")
    return docx_path

def read_input_path():
    try:
        with open("input_path.txt", "r") as f:
            input_path = f.read().strip()
        return input_path
    except FileNotFoundError:
        logging.error("input_path.txt not found")
        return None

input_path = read_input_path()
if not input_path:
    logging.error("No input path provided")
    sys.exit(1)

docx_path = get_docx_path(input_path)

converter = PDFConverter()

client = OpenAI()

load_dotenv()

app = QApplication(sys.argv)

converter.doc = win32.Dispatch("Word.Application").ActiveDocument
if converter.doc is None:
    logging.error("Failed to reference the already open DOCX file.")
    sys.exit(1)

converter.save_changes()
converter.save_changes()

pdf_path = converter.create_pdf_from_docx(docx_path)

if pdf_path:
    QMessageBox.information(None, 'Success', f'PDF created: {pdf_path}')
else:
    QMessageBox.warning(None, 'Error', 'Failed to create PDF from DOCX.')

def read_word_file(docx_path):
    doc = docx.Document(docx_path)
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
    return table_data

def get_gpt_response(user_input):
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are here to help extract data from tables."},
            {"role": "user", "content": user_input}
        ]
    )
    return response.choices[0].message.content.strip()

def extract_data_with_openai(table_data):
    prompt = (
        "Extract the descriptions and their corresponding amounts from the table with pricing. Exclude any rows where the description is a general statement or does not directly correspond to an amount. The descriptions should not be long or complicated.  You should only look for the simple ones, such as a city name and the state abreviation.  Short, brief descriptions.  If there are no brief descriptions, then just use what makes sense.  Also, in the amounts column, if there are 2 items, and the descriptions column has like a long description of cities and dates and stuff, but then there is also just city name and state abbreviations, and there are 2 of them? There you go, we need that.  Simple descriptions correlating with the amounts:\n"
        f"{table_data}\n"
        "Return only the list of tuples. Do not say anything else, just provide the list of tuples because your output is going to be read by a python script into a tuple. "
    )
    
    response = get_gpt_response(prompt)
    
    response_content = response.split('[')[-1].split(']')[0]
    response_content = '[' + response_content + ']'
    
    try:
        extracted_data = eval(response_content)
    except SyntaxError:
        raise ValueError("Failed to parse the response content as a list of tuples.")
    
    return extracted_data

def create_word_document(extracted_data, pdf_image_paths):
    new_doc = docx.Document()
    
    # Add the header part of the invoice
    header_lines = [
    os.getenv("HEADER_LINE_1"),
    os.getenv("HEADER_LINE_2"),
    os.getenv("HEADER_LINE_3"),
    os.getenv("HEADER_LINE_4")
    ]

    for line in header_lines:
        header_paragraph = new_doc.add_paragraph(line)
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header_paragraph.runs[0]
        header_run.font.size = Pt(11)
        header_run.font.name = 'Courier'
        header_paragraph.paragraph_format.line_spacing = 1
    
    new_doc.add_paragraph('')  # Add spacing after header

    # Add the rest of the invoice content for each row in extracted_data
    for job_id, (description, billing) in enumerate(extracted_data, start=1):
        description = description.upper()
        
        page_content = invoice.invoice_string.replace('<<billing>>', str(billing)) \
                                             .replace('<<description>>', description) \
                                             .replace('<<job>>', str(job_id))
        
        lines = page_content.split('\n')[5:]
        for line in lines:
            para = new_doc.add_paragraph(line)
            if "INVOICE NO." in line or "DATE:" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif "THANK YOU" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = 'Courier'
            para.paragraph_format.line_spacing = 1
        
        new_doc.add_page_break()
    
    # Handle PDF images if provided
    if pdf_image_paths:
    
        for image_path in pdf_image_paths:
            new_doc.add_picture(image_path)
            new_doc.add_page_break()
    
    new_doc.save('Formatted_Invoices.docx')
    print("Formatted document saved as Formatted_Invoices.docx")

def convert_pdf_to_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    image_paths = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        output_image_path = f"pdf_page_image_{page_num}.png"
        pix.save(output_image_path)
        image_paths.append(output_image_path)
    return image_paths

def main():
    word_file_path = docx_path
    if not word_file_path:
        print("No Word document selected. Exiting.")
        return
    
    pdf_file_path = pdf_path
    if not pdf_file_path:
        print("No PDF document selected. Exiting.")
        return

    table_data = read_word_file(word_file_path)
    extracted_data = extract_data_with_openai(table_data)

    app = QApplication(sys.argv)
    message = QMessageBox()
    message.setWindowTitle("Extracted Data")
    message.setText(f"Extracted Data:\n{extracted_data}")
    message.setIcon(QMessageBox.Information)
    message.exec_()

    
    df = pd.DataFrame(extracted_data, columns=["Description", "Amount"])

    # Expand rows by splitting "Amount" column where there are newline characters
    df = df.assign(Amount=df['Amount'].str.split('\n')).explode('Amount')

    # Clean up any unwanted characters (like $) and convert to float
    df['Amount'] = df['Amount'].replace('[\$,]', '', regex=True).astype(float)


    print(df)
    df.to_csv("extracted_data.csv", index=False)
    print("Data saved to extracted_data.csv")
    
    pdf_image_paths = convert_pdf_to_images(pdf_file_path)
    create_word_document(extracted_data, pdf_image_paths)

if __name__ == "__main__":
    main()







