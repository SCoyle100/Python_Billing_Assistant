from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL  # <-- IMPORT ADDED HERE

def split_large_amounts_and_format(input_path, output_path):
    # Load the document
    doc = Document(input_path)

    # Target the specific table based on manual inspection
    target_table = doc.tables[1]  # Adjust index if the table isn't the second one

    # Store the value from "Media Delivered = ..." for replacing the "Total" value
    media_delivered_value = None
    total_row = None
    rows_to_clear = []
    running_total = 0.0
    total_amount_indent = " " * 16

    # Process rows in the identified table
    rows_to_process = list(target_table.rows[8:])  # Start after header and metadata rows

    for idx, row in enumerate(rows_to_process):
        cells = row.cells
        description = cells[0].text.strip()  # First cell of each row for description
        amount_text = cells[3].text.strip()  # Fourth cell contains the amount

        # Check if this is the "Media Delivered = ..." row
        if "Media Delivered =" in description:
            try:
                media_delivered_value = float(
                    description.split('=')[-1]
                               .strip()
                               .replace(",", "")
                               .replace("$", "")
                )
                print(f"Extracted Media Delivered Value: ${media_delivered_value:,.2f}")  # Debug log
                rows_to_clear.append(row)  # Mark this row for clearing
            except ValueError:
                continue

        # Check if this is the "Total" row
        if "Total" in cells[3].text.strip():  # Identify "Total" in the 3rd column
            total_row = row
            # Log the current value in the Total row's 4th column for debugging
            total_value = cells[4].text.strip()
            print(f"Current Total Value: {total_value}")  # Debug log
            continue

        # Check for rows to clear (e.g., "Discount")
        if "discount" in description.lower():
            rows_to_clear.append(row)
            continue

        # Add dollar signs to all amounts and align them properly
        try:
            amount = float(amount_text.replace(",", "").replace("$", ""))
            if amount >= 1000:
                amount_indent = " " * 49
            else:
                amount_indent = " " * 52
            cells[3].text = f"{amount_indent}${amount:,.2f}"  # Add dollar sign with spacing
        except ValueError:
            continue

        # Add to our running total
        running_total += amount

        # Handle amounts greater than $5000 (splitting into parts)
        if amount > 5000:
            parts = []
            while amount > 0:
                part_amount = min(5000, amount)
                parts.append(part_amount)
                amount -= part_amount

            # Modify the description cell with parts
            indent = "      "  # Six spaces for city name and part labels
            new_description_lines = [f"{indent}{description}"] + [
                f"{indent}- PART {chr(64 + i)}" for i in range(1, len(parts) + 1)
            ]
            new_description_lines.append("")  # <-- Added blank line here
            cells[0].text = "\n".join(new_description_lines)

            # Modify the amount cell with parts
            new_amount_lines = [""]
            for part in parts:
                if part >= 1000:
                    part_indent = " " * 49
                else:
                    part_indent = " " * 52
                new_amount_lines.append(f"{part_indent}${part:,.2f}")
            new_amount_lines.append("")  # <-- Added blank line here
            cells[3].text = "\n".join(new_amount_lines)

        # Ensure font consistency for all text
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Match font size
                    run.font.name = "Times New Roman"  # Match font style

        # Replace the "Total" amount with the "Media Delivered" value
        if media_delivered_value is not None and total_row is not None:
            # Replace the value in cell[4] (5th column)
            total_row.cells[4].text = f"${media_delivered_value:,.2f}"  # Replace the value

    # Once done processing all rows, place the computed sum into the "Total" row
    if total_row is not None:
        total_row.cells[4].text = total_amount_indent + f"${running_total:,.2f}"

        # Now adjust paragraph alignment and spacing in the total amount cell
        for paragraph in total_row.cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

        # Keep "Total" text formatting as desired (e.g., Arial, size 16)
        for paragraph in total_row.cells[3].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(16)
                run.font.name = "Arial"

        # -------------------------
        # FINAL TOUCH: CENTER THE CELLS
        # -------------------------
        # Applies the "table properties → cell → center" alignment 
        # (vertical alignment + paragraph alignment).
        total_row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in total_row.cells[3].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        total_row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in total_row.cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Clear the contents of marked rows (Media Delivered and Discount rows)
    for row in rows_to_clear:
        for cell in row.cells:
            cell.text = ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"

    # Remove any fully empty rows
    rows_to_remove = []
    for row in target_table.rows:
        if all(cell.text.strip() == "" for cell in row.cells):
            rows_to_remove.append(row)

    for row in rows_to_remove:
        tbl = row._element.getparent()
        tbl.remove(row._element)

    # Ensure at least 2 rows remain. If so, add 2 blank lines to the second-to-last row.
    if len(target_table.rows) >= 2:
        second_to_last_row = target_table.rows[-2]
        # For example, add the blank lines in the first cell
        for _ in range(2):
            second_to_last_row.cells[0].add_paragraph("")

        # Optionally format the newly created blank lines
        for paragraph in second_to_last_row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

    # Save the modified document
    doc.save(output_path)
    print(f"Modified document saved as {output_path}")


# Define input and output file paths
input_file = r"D:\Programming\Billing_PDF_Automation\output\capitol_media.docx"
output_file = r"D:\Programming\Billing_PDF_Automation\output\capitol_media_updated.docx"

split_large_amounts_and_format(input_file, output_file)






















