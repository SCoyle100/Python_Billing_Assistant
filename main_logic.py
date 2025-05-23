import sys
import os
import re
import logging
import sqlite3
import datetime  # For generating batch IDs
from dotenv import load_dotenv
from PyQt5.QtWidgets import QApplication, QFileDialog, QInputDialog
from email import policy
from email.parser import BytesParser
import docx
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import dspy
import invoice  # Ensure your invoice template module is imported

from pdf_to_docx_ import PDFConverter

from database.database_functions import (
    save_invoices_to_db,
    BATCH_ID,

)

from vendor_invoice_logic.vendor_id import identify_vendors_from_pdfs_in_directory

from vendor_invoice_logic.matrix_media_logic import analyze_word_document

from vendor_invoice_logic.matrix_media_dataframe import (
    build_dataframe_from_word_document,
    
    
)

from vendor_invoice_logic.matrix_media_market_map import read_page_markets

from vendor_invoice_logic.capitol_media_logic import split_large_amounts_and_format


from vendor_invoice_logic.capitol_media_dataframe_1 import build_dataframe_from_capitol_media


from image_generation.create_pdf_image import create_images_from_docx

from utils.pdf_utils import combine_vendor_pdfs


#from vendor_invoice_logic.capitol_media_logic import split_large_amounts_and_format


# Global batch_id so that PDF and Email inserts share the same batch id within the same run.
#BATCH_ID = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")





converter = PDFConverter()


load_dotenv()

# Import performance decorators and logging config
from utils.decorators import performance_logger, cache_result, retry
from utils.logging_config import configure_logging

# Configure logging with timestamped files
configure_logging(logs_dir='logs', console_level=logging.INFO, file_level=logging.DEBUG)

# Initialize Qt Application for dialogs
app = QApplication(sys.argv)


# Configure DSPy with your OpenAI API key
dspy.configure(lm=dspy.LM('openai/gpt-4o'))




# Define DSPy signature with the user-provided date format
class ExtractInvoiceInfo(dspy.Signature):
    """
    Extract invoice information including description, amount, and job number if available. 
    Look for text like 'job:' or 'job #' or 'job number:' followed by an identifier (e.g., 'TTC-350').
    """
    
    text: str = dspy.InputField()
    invoices: list[dict[str, str]] = dspy.OutputField(
        desc="List of invoices, each with keys: Description, Amount, JobNumber (if available)"
    )


# Initialize the DSPy prediction module for extracting invoice info
invoice_extractor = dspy.Predict(ExtractInvoiceInfo)





import re

def format_job_number(job_number):
    """Format job numbers to ensure they have a hyphen between prefix and number."""
    if not job_number:
        return ""
        
    # If already has a hyphen, return as is
    if "-" in job_number:
        return job_number
        
    # Find the transition point between letters and numbers
    prefix = ""
    number = ""
    for i, char in enumerate(job_number):
        if char.isalpha() or char.isspace():
            prefix += char
        else:
            number = job_number[i:]
            break
            
    # Clean up prefix and number
    prefix = prefix.strip()
    number = number.strip()
    
    # If we found both parts, join with hyphen
    if prefix and number:
        return f"{prefix}-{number}"
    
    # Otherwise return original
    return job_number

def extract_job_number_from_description(description):
    """Extract job number from description text."""
    if not description:
        return "", description
        
    # Common job number patterns with capturing groups
    patterns = [
        # Format: "JOB: TTC-380" or "Job: TTC 380"
        r"\b(?:JOB|Job|job)\s*[:;#]\s*([A-Za-z]{2,4}[-\s]*\d{2,4})\b",
        
        # Format: "TTC-380" or "TTC 380" standalone
        r"\b([A-Za-z]{2,4}[-\s]*\d{2,4})\b",
        
        # Format: "Job #380" or simple numbers after job indicator
        r"\b(?:JOB|Job|job)\s*[:;#]\s*(\d{2,4})\b",
    ]
    
    job_number = ""
    clean_desc = description
    
    # Try each pattern until we find a match
    for pattern in patterns:
        match = re.search(pattern, clean_desc, re.IGNORECASE)
        if match:
            job_number = match.group(1)
            # Remove the entire match (not just the captured group)
            clean_desc = re.sub(pattern, "", clean_desc, flags=re.IGNORECASE)
            break
    
    # Also check for job number in parentheses
    if not job_number:
        # Look for patterns like "(JOB: TTC-380)" or "(Job #123)"
        parens_match = re.search(r"\(\s*(?:JOB|Job|job)\s*[:;#]?\s*([A-Za-z]{0,4}[-\s]*\d{2,4})\s*\)", clean_desc, re.IGNORECASE)
        if parens_match:
            job_number = parens_match.group(1)
            # Remove the entire parenthesized section
            clean_desc = re.sub(r"\(\s*(?:JOB|Job|job).*?\)", "", clean_desc, flags=re.IGNORECASE)
    
    # Also check for standalone "TTC-123" or similar patterns again
    if not job_number:
        standalone_match = re.search(r"\b([A-Za-z]{2,4}[-\s]*\d{2,4})\b", clean_desc)
        if standalone_match:
            job_number = standalone_match.group(1)
            # Only remove if it's clearly a job number and not part of a regular word
            if re.match(r"^[A-Za-z]{2,4}[-\s]*\d{2,4}$", job_number):
                clean_desc = re.sub(r"\b" + re.escape(job_number) + r"\b", "", clean_desc)
    
    # Clean up multiple spaces and trim
    clean_desc = re.sub(r'\s+', ' ', clean_desc).strip()
    
    return job_number, clean_desc

def clean_description_from_job_numbers(description):
    """Remove any job number references from the description."""
    if not description:
        return ""
    
    job_number, clean_desc = extract_job_number_from_description(description)
    return clean_desc

@performance_logger(output_dir='logs/performance')
def extract_structured_data_from_email(email_body):
    """
    Use DSPy to extract invoice information from the email body (description, amount, job_number).
    No duplicate checking is done; we simply return all extracted lines.
    """
    try:
        # Use DSPy to extract invoice information
        response = invoice_extractor(text=email_body)

        # Extract the list of invoices from the DSPy response
        structured_data = response.invoices

        # Process each invoice one by one to handle job number extraction and description cleaning
        extracted_data = []
        for invoice in structured_data:
            description = invoice.get("Description", "").upper()
            amount = str(invoice.get("Amount", "")).replace('$', '').replace(',', '')
            job_number = invoice.get("JobNumber", "")
            
            # Extract job number from description if not already provided
            extracted_job_number = ""
            if not job_number:
                extracted_job_number, description = extract_job_number_from_description(description)
                if extracted_job_number:
                    job_number = extracted_job_number
                    logging.info(f"Extracted job number '{job_number}' from description")
            else:
                # If a job number was already provided, still clean the description
                description = clean_description_from_job_numbers(description)
            
            # Format the job number with proper hyphen
            job_number = format_job_number(job_number)
            
            # Add the processed invoice data
            extracted_data.append((description, amount, job_number))
            
            # Log the extraction for debugging
            logging.info(f"Extracted: Description='{description}', Amount='{amount}', JobNumber='{job_number}'")
            
        # Log summary of extraction
        logging.info(f"Extracted {len(extracted_data)} invoices from email body")

        # Log the extracted data
        for index, data in enumerate(extracted_data):
            if len(data) >= 3 and data[2]:  # If job number is present
                logging.info(f"Extracted invoice #{index+1}: Description={data[0]}, Amount={data[1]}, JobNumber={data[2]}")
            else:
                logging.info(f"Extracted invoice #{index+1}: Description={data[0]}, Amount={data[1]}")

        return extracted_data
    except Exception as e:
        logging.error(f"Error during DSPy extraction: {e}")
        return None



def select_eml_file():
    options = QFileDialog.Options()
    options |= QFileDialog.ReadOnly
    file_path, _ = QFileDialog.getOpenFileName(None, "Select an EML File", "", 
                                               "Email Files (*.eml);;All Files (*)", 
                                               options=options)
    if file_path:
        process_selected_eml_file(file_path)
    else:
        logging.error("No EML file selected")

def process_selected_eml_file(eml_file_path):
    """
    Parse the selected EML file to extract email body content and attachments.
    If structured invoice data is found in the body, it will be inserted into
    the database first. Attachments (PDFs) are saved for subsequent processing.
    """
    logging.debug(f"Selected EML file: {eml_file_path}")

    # Parse the .eml file
    with open(eml_file_path, 'rb') as fp:
        msg = BytesParser(policy=policy.default).parse(fp)

    # Check and save attachments if any
    attachment_dir = os.path.join(os.getcwd(), 'downloaded files email')
    os.makedirs(attachment_dir, exist_ok=True)
    for part in msg.walk():
        content_disposition = part.get("Content-Disposition", "")
        if "attachment" in content_disposition:
            filename = part.get_filename()
            if filename:
                file_data = part.get_payload(decode=True)
                file_path = os.path.join(attachment_dir, filename)
                with open(file_path, 'wb') as f:
                    f.write(file_data)
                logging.info(f"Attachment {filename} saved to {attachment_dir}.")

    # Extract the plain text body of the email
    email_body = None
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == 'text/plain':
                email_body = part.get_content()
                break
    else:
        email_body = msg.get_content()

    if not email_body:
        logging.error("No plain text content found in the email.")
        return

    # Use DSPy to extract structured invoice data from the email content
    extracted_data = extract_structured_data_from_email(email_body)

    # If structured data is found, insert it into the DB before processing PDFs
    if extracted_data:
        save_invoices_to_db(
            invoices = extracted_data,
            batch_id = BATCH_ID,
            source = "FEE INVOICES"
        )

    else:
        logging.info("No structured data extracted from email body.")
            




@performance_logger(output_dir='logs/performance')
def process_all_pdfs_in_directory():
    """
    Loops through each PDF in 'downloaded files email' and calls handle_vendor_identification
    on a per-file basis. If multiple Matrix Media PDFs are found, they are combined into a single PDF
    before processing.
    """
    directory = "downloaded files email"
    
    # Identify vendors for all PDFs in the directory
    vendor_map = identify_vendors_from_pdfs_in_directory(directory)
    
    # If there are multiple Matrix Media PDFs, combine them into a single PDF
    matrix_media_files = [fname for fname, vendor in vendor_map.items() 
                         if vendor == "Matrix Media" and fname.lower().endswith(".pdf")]
    
    if len(matrix_media_files) > 1:
        logging.info(f"Found {len(matrix_media_files)} Matrix Media PDFs. Combining them into a single file.")
        
        # Sort the files alphabetically for consistent ordering
        matrix_media_files.sort()
        
        # Combine the Matrix Media PDFs
        combined_pdf_path = combine_vendor_pdfs(directory, "Matrix Media", vendor_map, "Combined_Matrix_Media.pdf")
        
        # Update the vendor map to include the new combined file
        if combined_pdf_path:
            new_filename = os.path.basename(combined_pdf_path)
            vendor_map[new_filename] = "Matrix Media"
            
            # Remove original files from vendor map as they've been combined
            for file in matrix_media_files:
                if file in vendor_map:
                    vendor_map.pop(file)
    
    # Get all PDF files in the directory
    all_pdf_files = [
        os.path.join(directory, f) for f in os.listdir(directory) 
        if os.path.isfile(os.path.join(directory, f)) and f.lower().endswith(".pdf")
    ]

    for pdf_file_path in all_pdf_files:
        # Skip original Matrix Media files if we created a combined file
        filename = os.path.basename(pdf_file_path)
        if len(matrix_media_files) > 1 and filename in matrix_media_files:
            logging.info(f"Skipping {filename} as it has been combined into a single PDF.")
            continue
            
        print(f"Processing file: {pdf_file_path}")
        handle_vendor_identification(pdf_file_path, vendor_map)


@performance_logger(output_dir='logs/performance')
def handle_vendor_identification(pdf_file_path, vendor_map=None):
    """
    Identifies the vendor for a single PDF file, then executes the appropriate logic.
    
    Args:
        pdf_file_path (str): Path to the PDF file to process.
        vendor_map (dict, optional): Mapping of filenames to vendor names. If None, 
                                    the function will generate it.
    """
    # If vendor_map is not provided, generate it for the current directory
    if vendor_map is None:
        vendor_map = identify_vendors_from_pdfs_in_directory(os.path.dirname(pdf_file_path))
    
    base_name = os.path.basename(pdf_file_path)
    vendor_name = vendor_map.get(base_name, "Unknown")

    print(f"{base_name} --> {vendor_name}")

    # Convert PDF to Word
    docx_file_path = converter.convert_pdf_to_docx(pdf_file_path)

    
    page_to_market = read_page_markets(docx_file_path)

    # Execute vendor-specific logic
    match vendor_name:
        case "Matrix Media":
            print(f"Executing script for {base_name}, vendor is Matrix Media...")
            # Apply the matrix media logic to update dollar amounts in the Word document
            analyze_word_document(docx_file_path)
            
            # Extract invoice data into a DataFrame
            df_invoices = build_dataframe_from_word_document(docx_file_path)
            
            # Debug print to verify DataFrame correctly identifies all markets
            print("DEBUG: DataFrame contents before converting to invoice list:")
            print(df_invoices)
            
            # Check if the DataFrame contains ServicePeriod and Description columns
            columns_to_include = ['Market', 'Amount']
            if 'ServicePeriod' in df_invoices.columns:
                columns_to_include.append('ServicePeriod')
            if 'Description' in df_invoices.columns:
                columns_to_include.append('Description')
                
            # Convert DataFrame rows to tuples with available columns
            invoices_list = list(df_invoices[columns_to_include].itertuples(index=False, name=None))
            
            print("DEBUG: Invoice list before saving to DB:")
            for invoice_tuple in invoices_list:
                if len(invoice_tuple) == 2:
                    market, amount = invoice_tuple
                    print(f"Market: '{market}', Amount: {amount}")
                elif len(invoice_tuple) == 3:
                    market, amount, service_period = invoice_tuple
                    print(f"Market: '{market}', Amount: {amount}, Service Period: '{service_period}'")
                elif len(invoice_tuple) == 4:
                    market, amount, service_period, description = invoice_tuple
                    print(f"Market: '{market}', Amount: {amount}, Service Period: '{service_period}', Description: '{description}'")
            
            # Save to database and get enhanced invoice data with invoice numbers
            enhanced_invoices = save_invoices_to_db(
                invoices=invoices_list,
                batch_id=BATCH_ID,
                source="Matrix Media",
                docx_file_path=docx_file_path  # Include the docx file path
            )

            print("DEBUG: Enhanced invoices after DB save:")
            for enhanced_invoice in enhanced_invoices:
                if len(enhanced_invoice) == 3:
                    market, amount, inv_no = enhanced_invoice
                    print(f"Market: '{market}', Amount: {amount}, Invoice: {inv_no}")
                elif len(enhanced_invoice) == 5:
                    market, amount, inv_no, service_period, description = enhanced_invoice
                    print(f"Market: '{market}', Amount: {amount}, Invoice: {inv_no}, Service Period: '{service_period}', Description: '{description}'")
            
            print("DEBUG: Page to market mapping:")
            for page, page_data in page_to_market.items():
                if isinstance(page_data, tuple) and len(page_data) == 2:
                    market, service_period = page_data
                    print(f"Page {page}: market='{market}', service_period='{service_period}'")
                else:
                    print(f"Page {page}: '{page_data}'")
                    
            # Ensure we use a simplified version of page_to_market with consistent service periods
            normalized_page_mapping = {}
            for page_num, page_data in page_to_market.items():
                if isinstance(page_data, tuple) and len(page_data) == 2:
                    market, service_period = page_data
                    # Normalize market name to match database
                    if any(fp in market.lower() for fp in ["fort payne", "ft. payne", "ft payne"]):
                        market = "Fort Payne"
                    normalized_page_mapping[page_num] = (market, service_period)
                else:
                    normalized_page_mapping[page_num] = (page_data, "")
            
            print("DEBUG: NORMALIZED Page to market mapping:")
            for page, page_data in normalized_page_mapping.items():
                market, service_period = page_data
                print(f"Page {page}: market='{market}', service_period='{service_period}'")

            # Create images from the Word document
            images = create_images_from_docx(
                docx_file_path, 
                "Matrix Media", 
                invoice_data=enhanced_invoices, 
                page_market_mapping=normalized_page_mapping
            )    


        case "Capitol Hill Media":
            print(f"Executing script for {base_name}, vendor is Capitol Hill Media...")
            split_large_amounts_and_format(docx_file_path)
            df_invoices = build_dataframe_from_capitol_media(docx_file_path)

            invoices_list = list(
            df_invoices[['Market', 'Amount']].itertuples(index=False, name=None)
             )
            

            '''
            save_invoices_to_db(
                invoices = invoices_list,
                batch_id = BATCH_ID,
                source = "Capitol Media"
                #docx_file_path = docx_file_path
            )

            images = create_images_from_docx(docx_file_path, vendor_name)
            if images:
                DOCX_IMAGES_MAP[docx_file_path] = images
                logging.info(f"Created {len(images)} images for {docx_file_path}.")
            else:
                logging.info(f"No images created for {docx_file_path}.")

            logging.debug(f"DOCX_IMAGES_MAP: {DOCX_IMAGES_MAP}")    

            '''

            enhanced_invoices = save_invoices_to_db(
                invoices=invoices_list,
                batch_id=BATCH_ID,
                source="Capitol Media",
                #docx_file_path=docx_file_path
            )
            images = create_images_from_docx(docx_file_path, vendor_name, enhanced_invoices, page_to_market)
            #if images:
            #    DOCX_IMAGES_MAP[docx_file_path] = images
            #    logging.info(f"Created {len(images)} images for {docx_file_path}.")

            


            #split_large_amounts_and_format()
            # call_capitol_hill_media_script(docx_file_path)  # your specialized logic
        case _:
            print(f"No specific handler for vendor: {vendor_name}")







from collections import defaultdict
import fnmatch
import glob



@performance_logger(output_dir='logs/performance')
def create_word_document():
    # Database fetch and filtering
    db_dir = os.path.join(os.getcwd(), 'database')
    db_path = os.path.join(db_dir, 'invoice.db')
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Properly log database connection to verify it's working
    logging.info(f"Connecting to database: {db_path}")
    if not os.path.exists(db_path):
        logging.error(f"Database file doesn't exist: {db_path}")
        return
        
    cursor.execute("""
        SELECT invoice_no, market, amount, batch_id, vendor, docx_file_path, service_period, description, job_number
        FROM invoices
        ORDER BY id  -- Ensure rows are ordered by insertion time
    """)
    all_rows = cursor.fetchall()
    conn.close()
    
    logging.info(f"Retrieved {len(all_rows)} rows from database")
    if not all_rows:
        logging.warning("No invoice data found in database")
        return

    # Group by batch_id to maintain the exact order of processing
    # No time filtering - just group all rows by batch_id
    batch_invoices = defaultdict(list)
    for row in all_rows:
        # Unpack row, handling old schema, newer schema, and newest schema with job_number
        if len(row) >= 9:
            invoice_no, market, amount, batch_id, source, docx_file_path, service_period, description, job_number = row
        elif len(row) >= 8:
            invoice_no, market, amount, batch_id, source, docx_file_path, service_period, description = row
            job_number = ""
        elif len(row) >= 7:
            invoice_no, market, amount, batch_id, source, docx_file_path, service_period = row
            description = ""
            job_number = ""
        else:
            invoice_no, market, amount, batch_id, source, docx_file_path = row
            service_period = ""
            description = ""
            job_number = ""
            
        try:
            # Just validate the batch_id format, don't filter by time
            datetime.datetime.strptime(batch_id, "%Y%m%d_%H%M%S")
            batch_invoices[batch_id].append(row)
        except ValueError:
            logging.warning(f"Invalid batch_id format: {batch_id}")
            continue
    
    # Get the most recent batch_id (we usually want to work with the latest batch)
    if not batch_invoices:
        logging.warning("No recent invoice data found")
        return
        
    latest_batch = sorted(batch_invoices.keys())[-1]
    logging.info(f"Processing latest batch: {latest_batch}")
    filtered_rows = batch_invoices[latest_batch]
    
    # Group invoices by source (aka vendor) maintaining the original order
    grouped_invoices = defaultdict(list)
    for row in filtered_rows:
        # Unpack row, handling old schema, newer schema, and newest schema with job_number
        if len(row) >= 9:
            invoice_no, market, amount, batch_id, source, docx_file_path, service_period, description, job_number = row
        elif len(row) >= 8:
            invoice_no, market, amount, batch_id, source, docx_file_path, service_period, description = row
            job_number = ""
        elif len(row) >= 7:
            invoice_no, market, amount, batch_id, source, docx_file_path, service_period = row
            description = ""
            job_number = ""
        else:
            invoice_no, market, amount, batch_id, source, docx_file_path = row
            service_period = ""
            description = ""
            job_number = ""
            
        # Include service_period, description, and job_number in the grouped invoices
        grouped_invoices[source].append((invoice_no, market, amount, batch_id, docx_file_path, service_period, description, job_number))

    logging.info(f"Grouped Invoices by vendor: {dict([(k, len(v)) for k, v in grouped_invoices.items()])}")

    # Initialize document
    new_doc = docx.Document()
    output_dir = os.path.join(os.getcwd(), 'final invoice output')
    os.makedirs(output_dir, exist_ok=True)
    control_chars_re = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]')

    def remove_control_characters(text):
        return control_chars_re.sub('', text)

    def add_invoice_page(doc, invoice_no, market, amount, add_pagebreak=True, description="", service_period="", job_number=""):
        """Add an invoice page with optional page break"""
        header_lines = [
            os.getenv("HEADER_LINE_1", ""),
            os.getenv("HEADER_LINE_2", ""),
            os.getenv("HEADER_LINE_3", ""),
            os.getenv("HEADER_LINE_4", "")
        ]
        for line in header_lines:
            header_paragraph = doc.add_paragraph(line)
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = header_paragraph.runs[0]
            header_run.font.size = Pt(11)
            header_run.font.name = 'Courier'
            header_paragraph.paragraph_format.line_spacing = 1

        doc.add_paragraph('')
        page_content = invoice.invoice_string  # from your "invoice" module
        page_content = page_content.replace('<<invoice>>', str(invoice_no))
        
        # Replace job number placeholder if available
        page_content = page_content.replace('<<job>>', str(job_number) if job_number else "")
        
        # Format description to start with market name
        # If we have both market and description, format as "Market - Description"
        # If service period is available, append it in parentheses
        if description and description.strip() and market and market.strip():
            # If description doesn't already start with the market name
            if not description.strip().startswith(market.strip()):
                display_text = f"{market} - {description}"
            else:
                display_text = description
        else:
            # Use whichever one is available (usually market)
            display_text = str(description) if description and description.strip() else str(market)
        
        # Add service period in parentheses if available
        if service_period and service_period.strip():
            display_text = f"{display_text} ({service_period})"
        
        # IMPORTANT: Don't append job number to description - it's handled separately in <<job>> placeholder
            
        page_content = page_content.replace('<<description>>', display_text)
        
        # Format the amount with dollar sign and two decimal places
        if isinstance(amount, str) and amount.startswith('$'):
            # If amount is already formatted with $, use it as is
            formatted_amount = amount
        else:
            # Otherwise, format it properly
            try:
                # Try to convert to float first (handles both string and numeric inputs)
                amount_float = float(amount)
                formatted_amount = f"${amount_float:.2f}"
            except (ValueError, TypeError):
                # If conversion fails, use as is
                formatted_amount = str(amount)
        
        page_content = page_content.replace('<<billing>>', formatted_amount)
        
        lines = page_content.split('\n')[5:]
        for line in lines:
            sanitized_line = remove_control_characters(line)
            para = doc.add_paragraph(sanitized_line)
            if "INVOICE NO." in line or "DATE:" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif "THANK YOU" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = 'Courier'
                para.paragraph_format.line_spacing = 1

        if add_pagebreak:
            doc.add_page_break()

    # Improved function to find images with extensive logging
    def find_invoice_images(invoice_no, market, vendor_name):
        logging.info(f"===== Searching for images: invoice={invoice_no}, market={market}, vendor={vendor_name} =====")
        
        # Define all directories where images might be stored (add more if needed)
        image_directories = [
            os.path.join(os.getcwd(), "downloaded files email"),
            os.path.join(os.getcwd(), "pdf images"),
            os.path.join(os.getcwd(), "images"),
            os.path.join(os.getcwd(), "output"),
            os.getcwd()  # Check root directory too
        ]
        
        # Log directories we're searching
        logging.info(f"Searching in directories: {image_directories}")
        
        matching_images = []
        
        # Check if this is a Fort Payne invoice
        is_fort_payne = False
        if market and any(fp in market.lower() for fp in ["fort payne", "ft. payne", "ft payne"]):
            is_fort_payne = True
            logging.info(f"This is a Fort Payne invoice: {invoice_no}")
        
        # Try several different patterns, from most specific to most general
        patterns = []
        
        # Normalize inputs for filename matching
        safe_invoice_no = "".join(c for c in str(invoice_no) if c.isalnum() or c in ('-', '_'))
        safe_market = "".join(c for c in str(market) if c.isalnum() or c in ('-', '_')).lower()
        safe_vendor = "".join(c for c in str(vendor_name) if c.isalnum() or c in ('-', '_')).lower()
        
        # Special patterns for Fort Payne
        if is_fort_payne and vendor_name == "Matrix Media":
            # For Fort Payne, we need to check for various spellings/formats
            patterns.append((f"{safe_invoice_no}_fortpayne_{safe_vendor}_page_*.png", "Fort Payne exact"))
            patterns.append((f"{safe_invoice_no}_fort*payne*_page_*.png", "Fort Payne wildcard"))
            patterns.append((f"{safe_invoice_no}_ft*payne*_page_*.png", "Ft Payne wildcard"))
        
        # Standard patterns
        # Pattern 1: Exact match with invoice, market, vendor
        patterns.append((f"{safe_invoice_no}_{safe_market}_{safe_vendor}_page_*.png", "exact match"))
        
        # Pattern 2: Just invoice number and page
        patterns.append((f"{safe_invoice_no}_*page_*.png", "invoice number with page"))
        
        # Pattern 3: Any file containing the invoice number
        patterns.append((f"*{safe_invoice_no}*.png", "contains invoice number"))
        
        # For each directory
        for image_dir in image_directories:
            if not os.path.exists(image_dir):
                logging.debug(f"Directory does not exist: {image_dir}")
                continue
                
            logging.info(f"Checking directory: {image_dir}")
            
            # List all PNG files in the directory for logging
            png_files = [f for f in os.listdir(image_dir) if f.lower().endswith('.png')]
            if png_files:
                logging.info(f"Found {len(png_files)} PNG files in {image_dir}")
                logging.debug(f"PNG files: {png_files[:10]}")  # List up to 10 PNG files for debugging
            
            # Try each pattern until we find matches
            for pattern, pattern_desc in patterns:
                logging.debug(f"Trying pattern: {pattern} ({pattern_desc})")
                pattern_matches = []
                
                for f in os.listdir(image_dir):
                    if f.lower().endswith('.png') and fnmatch.fnmatch(f.lower(), pattern.lower()):
                        image_path = os.path.join(image_dir, f)
                        pattern_matches.append(image_path)
                        
                if pattern_matches:
                    logging.info(f"Found {len(pattern_matches)} matches with pattern '{pattern_desc}'")
                    matching_images.extend(pattern_matches)
                    break  # Skip remaining patterns for this directory
        
        if not matching_images:
            logging.warning(f"⚠️ NO IMAGES FOUND for invoice {invoice_no}, market {market}, vendor {vendor_name}")
        else:
            logging.info(f"Found {len(matching_images)} total images: {[os.path.basename(img) for img in matching_images]}")
        
        # Sort images if we found any
        if matching_images:
            # Try to sort by page number if possible
            try:
                matching_images.sort(key=lambda x: int(os.path.basename(x).split('_page_')[1].split('.')[0]))
            except (IndexError, ValueError):
                # If can't sort by page number, sort by filename
                matching_images.sort()
                
        return matching_images

    # Process vendors in a specific order based on your requirements
    vendor_processing_order = ["FEE INVOICES", "Matrix Media", "Capitol Media"]
    
    # Keep track of images that have been inserted to avoid duplicates
    processed_images = set()
    
    # Keep track of invoice numbers that have been processed to avoid duplicate image insertions
    processed_invoice_numbers = set()
    
    # Debug counter for image insertions
    image_insert_count = 0
    
    # Process each vendor in the desired order
    for vendor_name in vendor_processing_order:
        if vendor_name not in grouped_invoices:
            logging.info(f"No invoices found for vendor: {vendor_name}")
            continue
            
        invoice_list = grouped_invoices[vendor_name]
        logging.info(f"Processing vendor/source: {vendor_name} with {len(invoice_list)} invoices")
        
        # For Capitol Media, we'll collect all invoice images to add after all invoices
        capitol_media_all_images = []

        # Build each invoice page
        for invoice_data in invoice_list:
            # Unpack invoice data with variable length handling
            if len(invoice_data) >= 8:
                invoice_no, market, amount, batch_id, docx_file_path, service_period, description, job_number = invoice_data
            elif len(invoice_data) >= 7:
                invoice_no, market, amount, batch_id, docx_file_path, service_period, description = invoice_data
                job_number = ""
            elif len(invoice_data) >= 6:
                invoice_no, market, amount, batch_id, docx_file_path, service_period = invoice_data
                description = ""
                job_number = ""
            else:
                invoice_no, market, amount, batch_id, docx_file_path = invoice_data
                service_period = ""
                description = ""
                job_number = ""
                
            log_msg = f"Adding invoice: {invoice_no}, market: {market}, amount: {amount}, service_period: {service_period}"
            if job_number:
                log_msg += f", job: {job_number}"
            logging.info(log_msg)
            
            # Find matching images for this invoice (do this only once)
            # Use a composite key that includes service period to handle duplicate markets with different service periods
            invoice_key = f"{invoice_no}_{market}_{service_period}".replace(" ", "_").lower()
            
            # If we've already processed this specific invoice for this market and service period, skip it
            if invoice_key in processed_invoice_numbers:
                logging.info(f"Skipping already processed invoice: {invoice_no}, market: {market}, service period: {service_period}")
                continue
                
            # Mark this invoice as processed
            processed_invoice_numbers.add(invoice_key)
            
            # Initialize image cache if needed
            if not hasattr(create_word_document, 'image_cache'):
                create_word_document.image_cache = {}
                
            # Try to get images from cache or find them
            if invoice_key in create_word_document.image_cache:
                matching_images = create_word_document.image_cache[invoice_key]
                logging.info(f"Using cached images for {invoice_no}, {market}, {service_period}")
            else:
                matching_images = find_invoice_images(invoice_no, market, vendor_name)
                
                # Special handling for Fort Payne if no images found in the regular search
                is_fort_payne = vendor_name == "Matrix Media" and (
                    "Fort Payne" in market or "Ft. Payne" in market or "Ft Payne" in market
                )
                
                if is_fort_payne and not matching_images:
                    logging.info(f"Fort Payne invoice with no images - searching for any Fort Payne images")
                    # Build a more general pattern for Fort Payne
                    fort_payne_pattern = f"*{invoice_no}*fort*payne*.png"
                    fort_payne_images = []
                    
                    # Search in all image directories
                    for image_dir in [
                        os.path.join(os.getcwd(), "downloaded files email"),
                        os.path.join(os.getcwd(), "pdf images"),
                        os.path.join(os.getcwd(), "images"),
                        os.path.join(os.getcwd(), "output"),
                        os.getcwd()
                    ]:
                        if os.path.exists(image_dir):
                            for f in os.listdir(image_dir):
                                if f.lower().endswith('.png') and fnmatch.fnmatch(f.lower(), fort_payne_pattern.lower()):
                                    fort_payne_images.append(os.path.join(image_dir, f))
                    
                    if fort_payne_images:
                        logging.info(f"Found {len(fort_payne_images)} Fort Payne images for invoice {invoice_no}")
                        matching_images = fort_payne_images
                
                # Cache the images we found (including Fort Payne special search results)
                create_word_document.image_cache[invoice_key] = matching_images
                
            has_images = len(matching_images) > 0
            
            # Add the invoice page with description, service period, and job number
            add_invoice_page(
                new_doc, 
                invoice_no, 
                market, 
                amount, 
                not has_images,  # Only add page break if no images
                description=description,
                service_period=service_period,
                job_number=job_number
            )
            
            # Handle images based on vendor type
            if vendor_name in ["Matrix Media", "FEE INVOICES"]:
                # For both Matrix Media and FEE INVOICES, add images directly after the invoice
                logging.info(f"Processing images for invoice_key: {invoice_key}")
                
                if matching_images:
                    logging.info(f"Adding {len(matching_images)} images for {vendor_name} invoice {invoice_no} - market: '{market}', service period: '{service_period}'")
                    images_added = 0
                    for img_path in matching_images:
                        # Skip images we've already processed
                        if img_path in processed_images:
                            logging.info(f"Skipping already processed image: {img_path}")
                            continue
                            
                        try:
                            logging.info(f"Adding image to document: {img_path}")
                            new_doc.add_page_break()
                            new_doc.add_picture(img_path, width=Inches(6))
                            processed_images.add(img_path)  # Mark as processed
                            images_added += 1
                            image_insert_count += 1
                            logging.info(f"Successfully added image: {img_path} (Total images: {image_insert_count})")
                        except Exception as e:
                            logging.error(f"Error adding image {img_path}: {str(e)}")
                    
                    # Only add a page break if we actually added images
                    if images_added > 0:
                        new_doc.add_page_break()
                else:
                    logging.warning(f"No images found for {vendor_name} invoice {invoice_no}")
            
            elif vendor_name == "Capitol Media":
                # For Capitol Media, collect all images to add after all invoices
                if matching_images:
                    logging.info(f"Collecting {len(matching_images)} images for Capitol Media invoice {invoice_no}")
                    capitol_media_all_images.extend(matching_images)
        
        # For Capitol Media, add all collected images after all invoices
        # NOTE: We implemented duplicate prevention for Matrix Media above by tracking invoice_numbers.
        # Similar changes might be needed here for Capitol Media if duplicate images are observed.
        if vendor_name == "Capitol Media" and capitol_media_all_images:
            logging.info(f"Adding {len(capitol_media_all_images)} images for all Capitol Media invoices")
            for img_path in capitol_media_all_images:
                try:
                    logging.info(f"Adding image to document: {img_path}")
                    new_doc.add_page_break()
                    new_doc.add_picture(img_path, width=Inches(6))
                    logging.info(f"Successfully added image: {img_path}")
                except Exception as e:
                    logging.error(f"Error adding image {img_path}: {str(e)}")
            new_doc.add_page_break()
        elif vendor_name == "Capitol Media" and not capitol_media_all_images:
            logging.warning(f"No images found for Capitol Media vendor")

    # Save the assembled Word doc with the batch ID in the filename
    output_path = os.path.join(output_dir, f'final_invoice_output_{latest_batch}.docx')
    try:
        new_doc.save(output_path)
        logging.info(f"Formatted document saved as {output_path}")
        
        # Also save a copy with a generic name for easy access
        standard_output_path = os.path.join(output_dir, 'final_invoice_output.docx')
        new_doc.save(standard_output_path)
        logging.info(f"Formatted document also saved as {standard_output_path}")
    except Exception as e:
        logging.error(f"Error saving document: {str(e)}")
    
    return output_path



if __name__ == "__main__":
    select_eml_file()
    process_all_pdfs_in_directory()
    create_word_document()