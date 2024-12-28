from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def split_large_amounts_and_format(input_path, output_path):
    # Load the document
    doc = Document(input_path)

    # Target the specific table based on manual inspection
    target_table = doc.tables[1]  # Adjust index if the table isn't the second one

    # Store the value from "Media Delivered = ..." for replacing the "Total" value
    media_delivered_value = None
    total_row = None
    rows_to_clear = []

    # Process rows in the identified table
    rows_to_process = list(target_table.rows[8:])  # Start after header and metadata rows

    for idx, row in enumerate(rows_to_process):
        cells = row.cells
        description = cells[0].text.strip()  # First cell of each row for description
        amount_text = cells[3].text.strip()  # Fourth cell contains the amount

        # Check if this is the "Media Delivered = ..." row
        if "Media Delivered =" in description:
            try:
                media_delivered_value = float(description.split('=')[-1].strip().replace(",", "").replace("$", ""))
                print(f"Extracted Media Delivered Value: ${media_delivered_value:,.2f}")  # Debug log
                rows_to_clear.append(row)  # Mark this row for clearing
            except ValueError:
                continue

        # Check if this is the "Total" row
        if "Total" in cells[3].text.strip():  # Identify "Total" in the 3rd column
            total_row = row
            # Log the current value in the Total row's 4th column for debugging
            total_value = cells[4].text.strip()
            print(f"Current Total Value: {total_value}")  # Debug log
            continue

        # Check for rows to clear (e.g., "Discount")
        if "discount" in description.lower():
            rows_to_clear.append(row)
            continue

        # Add dollar signs to all amounts and align them properly
        try:
            amount = float(amount_text.replace(",", "").replace("$", ""))
            if amount >= 1000:
                amount_indent = " " * 49
            else:
                amount_indent = " " * 52
            cells[3].text = f"{amount_indent}${amount:,.2f}"  # Add dollar sign with spacing
        except ValueError:
            continue

        # Handle amounts greater than $5000 (splitting into parts)
        if amount > 5000:
            parts = []
            while amount > 0:
                part_amount = min(5000, amount)
                parts.append(part_amount)
                amount -= part_amount

            # Modify the description cell with parts
            indent = "      "  # Six spaces for city name and part labels
            new_description_lines = [f"{indent}{description}"] + [
                f"{indent}- PART {chr(64 + i)}" for i in range(1, len(parts) + 1)
            ]
            cells[0].text = "\n".join(new_description_lines)

            # Modify the amount cell with parts
            new_amount_lines = [""]
            for part in parts:
                if part >= 1000:
                    part_indent = " " * 49
                else:
                    part_indent = " " * 52
                new_amount_lines.append(f"{part_indent}${part:,.2f}")
            cells[3].text = "\n".join(new_amount_lines)

        # Ensure font consistency for all text
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Match font size
                    run.font.name = "Times New Roman"  # Match font style

        # Replace the "Total" amount with the "Media Delivered" value
        if media_delivered_value is not None and total_row is not None:
            # Replace the value in cell[4] (5th column)
            total_row.cells[4].text = f"${media_delivered_value:,.2f}"  # Replace the value
    for paragraph in total_row.cells[4].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align text
        for run in paragraph.runs:
            run.font.size = Pt(9)  # Set font size to 11
            run.font.name = "Times New Roman"  # Set font to Arial

    # Keep "Total" text formatting as Arial, size 11
    for paragraph in total_row.cells[3].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)  # Set font size to 11
            run.font.name = "Arial"  # Set font to Arial


    # Clear the contents of marked rows (Media Delivered and Discount rows)
    for row in rows_to_clear:
        for cell in row.cells:
            cell.text = ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"

    # Save the modified document
    doc.save(output_path)
    print(f"Modified document saved as {output_path}")

# Define input and output file paths
input_file = "D:\Programming\Billing_PDF_Automation\output\Thompson Tractor-4Q #1835 November 2024 Media Reconciliation Invoice_modified.docx"  # Replace with the path to your input file
output_file = "D:\Programming\Billing_PDF_Automation\output\Thompson Tractor-4Q #1835 November 2024 Media Reconciliation Invoice_modified_final.docx"  # Replace with the desired output file path

# Apply the function
split_large_amounts_and_format(input_file, output_file)





















