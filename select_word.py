import sys
from PyQt5.QtWidgets import QApplication, QMessageBox, QFileDialog
import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import os
from dotenv import load_dotenv
import invoice
import fitz  # PyMuPDF
import logging
import win32com.client as win32
from pdf_to_docx import PDFConverter



converter = PDFConverter()


client = OpenAI()

load_dotenv()

logging.basicConfig(level=logging.DEBUG)

# Function to prompt the user to select a Word document
def select_word_document():
    app = QApplication(sys.argv)
    options = QFileDialog.Options()
    options |= QFileDialog.ReadOnly
    file_path, _ = QFileDialog.getOpenFileName(None, "Select a Word Document", "", "Word Files (*.docx);;All Files (*)", options=options)
    if file_path:
        process_selected_word_document(file_path)
    else:
        logging.error("No Word document selected")

        

'''
def read_word_file(file_path):
    doc = docx.Document(file_path)
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
    return table_data
    '''

def read_word_document_as_string(file_path):
    doc = docx.Document(file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text.strip())

    return "\n".join(full_text)



# Function to process the selected Word document
def process_selected_word_document(file_path):
    logging.debug(f"Selected Word file: {file_path}")
    
    # Assuming the rest of the code works with the selected Word document
    table_data = read_word_document_as_string(file_path)
    extracted_data = extract_data_with_openai(table_data)

    app = QApplication(sys.argv)
    message = QMessageBox()
    message.setWindowTitle("Extracted Data")
    message.setText(f"Extracted Data:\n{table_data}")
    message.setIcon(QMessageBox.Information)
    message.exec_()
    
    df = pd.DataFrame(extracted_data, columns=["Description", "Amount"])
    
    # Remove non-numeric characters and convert to numeric
    df['Amount'] = df['Amount'].replace('[\$,]', '', regex=True).astype(float)

    print(df)
    df.to_csv("extracted_data.csv", index=False)
    print("Data saved to extracted_data.csv")

    pdf_path = converter.create_pdf_from_docx(file_path)

    if pdf_path:
        QMessageBox.information(None, 'Success', f'PDF created: {pdf_path}')
    else:
        QMessageBox.warning(None, 'Error', 'Failed to create PDF from DOCX.')
    
    # Assuming there's a need to create a new Word document
    create_word_document(extracted_data, [])





def get_gpt_response(user_input):
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are here to help extract data from tables."},
            {"role": "user", "content": user_input}
        ]
    )
    return response.choices[0].message.content.strip()

def extract_data_with_openai(table_data):
    prompt = (
        "Extract the description and amounts column from the main table.  Start the search starting at the word Description:\n"
        f"{table_data}\n"
        "Return only the list of tuples."
    )
    
    response = get_gpt_response(prompt)
    
    # Clean up the response to ensure it's valid Python code for a list of tuples
    response_content = response.split('[')[-1].split(']')[0]
    response_content = '[' + response_content + ']'
    
    try:
        extracted_data = eval(response_content)
    except SyntaxError:
        raise ValueError("Failed to parse the response content as a list of tuples.")
    
    return extracted_data
def create_word_document(extracted_data, pdf_image_paths):
    new_doc = docx.Document()
    
    # Add the rest of the invoice content for each row in extracted_data
    for job_id, (description, billing) in enumerate(extracted_data, start=1):
        # Convert description to uppercase
        description = description.upper()
        
        # Add and format the header part of the invoice_string
        header_lines = [
        os.getenv("HEADER_LINE_1"),
        os.getenv("HEADER_LINE_2"),
        os.getenv("HEADER_LINE_3"),
        os.getenv("HEADER_LINE_4")
        ]
        
        for line in header_lines:
            header_paragraph = new_doc.add_paragraph(line)
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment for header
            header_run = header_paragraph.runs[0]
            header_run.font.size = Pt(11)  # Set font size to 11 for the header
            header_run.font.name = 'Courier'  # Set font to Courier
            
            # Set line spacing to single
            header_paragraph.paragraph_format.line_spacing = 1
        
        # Add an empty paragraph for spacing between header and content
        new_doc.add_paragraph('')

        # Replace placeholders with data
        page_content = invoice.invoice_string.replace('<<billing>>', str(billing)) \
                                             .replace('<<description>>', description) \
                                             .replace('<<job>>', str(job_id))
        
        # Skip the header part that was already added and process the rest of the content
        lines = page_content.split('\n')[5:]
        for line in lines:
            para = new_doc.add_paragraph(line)
            
            # Align "INVOICE NO.: <<invoice>>" and "DATE: <<date>>" to the right
            if "INVOICE NO.:" in line or "DATE:" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right alignment for invoice and date
            elif "THANK YOU" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment for THANK YOU
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left alignment for other content
            
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)  # Set font size to 9 for the content
                run.font.name = 'Courier'  # Set font to Courier
            
            # Set line spacing to single
            para.paragraph_format.line_spacing = 1
        
        # Add a page break after processing each row's content
        new_doc.add_page_break()
    
    # Handle PDF images if provided
    if pdf_image_paths:
        new_doc.add_paragraph("Attached Images from PDF:")
        for image_path in pdf_image_paths:
            new_doc.add_picture(image_path)
            new_doc.add_page_break()
    
    # Save the document
    new_doc.save('Formatted_Invoices.docx')
    print("Formatted document saved as Formatted_Invoices.docx")


def convert_pdf_to_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    image_paths = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        output_image_path = f"pdf_page_image_{page_num}.png"
        pix.save(output_image_path)
        image_paths.append(output_image_path)
    return image_paths

# Assuming this is the original entry point to the script
if __name__ == "__main__":
    select_word_document()
