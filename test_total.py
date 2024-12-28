from docx import Document
from docx.shared import Pt

def split_large_amounts_and_format(input_path, output_path):
    # Load the document
    doc = Document(input_path)

    # Target the specific table based on manual inspection
    target_table = doc.tables[1]  # Adjust index if the table isn't the second one

    # Store the value from "Media Delivered = ..." for replacing the "Total" value
    media_delivered_value = None
    total_row = None

    # Process rows in the identified table
    for idx, row in enumerate(target_table.rows):
        cells = row.cells
        # Ensure there are at least 5 cells in the row
        if len(cells) < 5:
            continue

        # Debugging: Log the text in the columns for each row
        col_4_text = cells[3].text.strip().lower() if len(cells) > 3 else "(Empty)"
        col_5_text = cells[4].text.strip() if len(cells) > 4 else "(Empty)"
        print(f"Row {idx}: Col 4: '{col_4_text}', Col 5: '{col_5_text}'")

        # Check for "Media Delivered = ..." in the third column
        if "media delivered =" in col_4_text:
            try:
                media_delivered_value = float(cells[3].text.strip().split('=')[-1].replace(",", "").replace("$", ""))
                print(f"Captured Media Delivered Value: ${media_delivered_value:,.2f}")  # Debug log
            except ValueError:
                print(f"Error parsing Media Delivered value in row {idx}.")
                continue

        # Check for "Total" in the fifth column
        if col_5_text.lower() == "total":
            total_row = row
            total_value = cells[4].text.strip()  # Log the current value in the fifth column
            print(f"Captured Total Row at {idx}: Current Total Value: {total_value}")

    # Replace the "Total" value with the "Media Delivered" value
    if media_delivered_value is not None and total_row is not None:
        total_row.cells[4].text = f"${media_delivered_value:,.2f}"  # Replace the value
        for paragraph in total_row.cells[4].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
        print(f"Updated Total Value to: ${media_delivered_value:,.2f}")

    else:
        print("Total row or Media Delivered value not found. Please check the document structure.")

    # Save the modified document
    doc.save(output_path)
    print(f"Modified document saved as {output_path}")

# Define input and output file paths
input_file = "D:\Programming\Billing_PDF_Automation\output\Thompson Tractor-4Q #1835 November 2024 Media Reconciliation Invoice_modified.docx"  # Replace with the path to your input file
output_file = "D:\Programming\Billing_PDF_Automation\output\Thompson Tractor-4Q #1835 November 2024 Media Reconciliation Invoice_modified_final.docx"  # Replace with the desired output file path

# Apply the function
split_large_amounts_and_format(input_file, output_file)


