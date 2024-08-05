import os
import logging
from xml.etree.ElementTree import QName
from dotenv import load_dotenv
import sys
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
from adobe.pdfservices.operation.exception.exceptions import ServiceApiException, ServiceUsageException, SdkException
from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
from adobe.pdfservices.operation.io.stream_asset import StreamAsset
from adobe.pdfservices.operation.pdf_services import PDFServices
from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult
from adobe.pdfservices.operation.pdfjobs.jobs.create_pdf_job import CreatePDFJob
from adobe.pdfservices.operation.pdfjobs.result.create_pdf_result import CreatePDFResult
from openai import OpenAI
from PyQt5.QtWidgets import QFileDialog, QApplication, QMessageBox
import win32com.client as win32

logging.basicConfig(level=logging.INFO)

class PDFConverter:
    def __init__(self):
        self.word = None
        self.doc = None

    def convert_pdf_to_docx(self, input_path):
        try:
            with open(input_path, 'rb') as file:
                input_stream = file.read()

            load_dotenv()

            credentials = ServicePrincipalCredentials(
                client_id=os.getenv('PDF_SERVICES_CLIENT_ID'),
                client_secret=os.getenv('PDF_SERVICES_CLIENT_SECRET')
            )

            pdf_services = PDFServices(credentials=credentials)
            input_asset = pdf_services.upload(input_stream=input_stream, mime_type=PDFServicesMediaType.PDF)
            export_pdf_params = ExportPDFParams(target_format=ExportPDFTargetFormat.DOCX)
            export_pdf_job = ExportPDFJob(input_asset=input_asset, export_pdf_params=export_pdf_params)
            location = pdf_services.submit(export_pdf_job)
            pdf_services_response = pdf_services.get_job_result(location, ExportPDFResult)
            result_asset = pdf_services_response.get_result().get_asset()
            stream_asset = pdf_services.get_content(result_asset)

            output_file_path = self.create_output_file_path(input_path)
            with open(output_file_path, "wb") as file:
                file.write(stream_asset.get_input_stream())

            return output_file_path

        except (ServiceApiException, ServiceUsageException, SdkException) as e:
            logging.exception(f'Exception encountered while executing operation: {e}')
            return None

    def create_output_file_path(self, input_path):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(script_dir, "output")
        os.makedirs(output_dir, exist_ok=True)

        base_name = os.path.basename(input_path)
        file_name, _ = os.path.splitext(base_name)
        output_file_path = os.path.join(output_dir, f"{file_name}.docx")

        return output_file_path

    def open_and_edit_docx(self, file_path):
        self.word = win32.gencache.EnsureDispatch('Word.Application')
        self.word.Visible = True
        self.doc = self.word.Documents.Open(file_path)

    def save_changes(self):
        self.doc.Save()

    def create_pdf_from_docx(self, docx_path):
        try:
            with open(docx_path, 'rb') as file:
                input_stream = file.read()

            load_dotenv()

            credentials = ServicePrincipalCredentials(
                client_id=os.getenv('PDF_SERVICES_CLIENT_ID'),
                client_secret=os.getenv('PDF_SERVICES_CLIENT_SECRET')
            )

            pdf_services = PDFServices(credentials=credentials)
            input_asset = pdf_services.upload(input_stream=input_stream, mime_type=PDFServicesMediaType.DOCX)
            create_pdf_job = CreatePDFJob(input_asset)
            location = pdf_services.submit(create_pdf_job)
            pdf_services_response = pdf_services.get_job_result(location, CreatePDFResult)
            result_asset = pdf_services_response.get_result().get_asset()
            stream_asset = pdf_services.get_content(result_asset)

            output_file_path = self.create_pdf_output_file_path(docx_path)
            with open(output_file_path, "wb") as file:
                file.write(stream_asset.get_input_stream())

            return output_file_path

        except (ServiceApiException, ServiceUsageException, SdkException) as e:
            logging.exception(f'Exception encountered while executing operation: {e}')
            return None

    def create_pdf_output_file_path(self, input_path):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(script_dir, "output")
        os.makedirs(output_dir, exist_ok=True)

        base_name = os.path.basename(input_path)
        file_name, _ = os.path.splitext(base_name)
        now = datetime.now()
        time_stamp = now.strftime("%Y-%m-%dT%H-%M-%S")
        output_file_path = os.path.join(output_dir, f"{file_name}_{time_stamp}.pdf")

        return output_file_path

def read_word_file(docx_path):
    doc = docx.Document(docx_path)
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
    return table_data

def get_gpt_response(user_input):
    client = OpenAI()
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are here to help extract data from tables."},
            {"role": "user", "content": user_input}
        ]
    )
    return response.choices[0].message.content.strip()

def extract_data_with_openai(table_data):
    prompt = (
        "Extract the amounts from the table. Ignore the descriptions. Just provide the list of amounts.  Ignore any amounts with a negative amount, and ignore any amounts that are preceded by an equals sign. And do not include any Totals.  Do not repeat any amounts.:\n"
        f"{table_data}\n"
        "Return only the list of amounts as a Python list. Do not say anything else."
    )
    response = get_gpt_response(prompt)
    response = response.strip().strip('```python').strip('```')
    
    try:
        extracted_data = eval(response)
    except SyntaxError:
        raise ValueError("Failed to parse the response content as a list.")
    
    return extracted_data

def split_large_amounts(amounts):
    new_amounts = []
    for amount in amounts:
        # If the amount is a string, remove commas
        if isinstance(amount, str):
            amount = amount.replace(',', '')
            try:
                amount_float = float(amount)
            except ValueError:
                raise ValueError(f"Failed to convert amount to float: {amount}")
        else:
            # If it's already a float, use it directly
            amount_float = amount

        if amount_float >= 5000.00:
            split_amount1 = amount_float * 0.6
            split_amount2 = amount_float * 0.4
            
            # Adjust the split to ensure both are under 4900
            while split_amount1 > 4900 or split_amount2 > 4900:
                split_amount1 *= 0.9
                split_amount2 = amount_float - split_amount1

            new_amounts.extend([split_amount1, split_amount2])
        else:
            new_amounts.append(amount_float)
    
    # Debug: Print the final processed amounts
    print("Final Processed Amounts: ", new_amounts)

    return new_amounts





def set_cell_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')  # Border width
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), '000000')  # Border color
        tcBorders.append(border_el)
    
    tcPr.append(tcBorders)

def create_invoice_table(doc, extracted_data):
    # Insert a page break to move the sentence to the second page
    doc.add_page_break()
    
    doc.add_paragraph("Fill This Out")
    
    table = doc.add_table(rows=1, cols=3)

    # Set up the header row with bold and slightly larger text
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Amount'
    hdr_cells[2].text = 'Invoice Number'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        set_cell_border(cell)

    # Check if the extracted data is processed correctly
    print("Extracted and Split Amounts: ", extracted_data)

    # Add a row for each amount
    for amount in extracted_data:
        row_cells = table.add_row().cells
        row_cells[0].text = ''  # Add description if needed
        row_cells[1].text = f'{amount:.2f}'  # Format amount as a string with 2 decimal places
        row_cells[2].text = ''  # Add invoice number if needed
        for cell in row_cells:
            set_cell_border(cell)

    return table

def create_job_number_field(doc, previous_table):
    # Create a 1-column table below the previous table
    job_table = doc.add_table(rows=1, cols=1)
    job_table.rows[0].cells[0].text = 'Job Number:'
    
    # Set width of job number table to be half the width of the previous table
    previous_table_width = previous_table.columns[0].width
    job_table.columns[0].width = previous_table_width // 2
    
    for cell in job_table.rows[0].cells:
        set_cell_border(cell)


def save_and_close_initial_doc(converter):
    # Save and close the initial Word document
    converter.save_changes()
    converter.word.Quit()

def create_modified_doc_with_table(docx_path, extracted_data):
    # Open the saved DOCX file
    doc = docx.Document(docx_path)

    # Split large amounts
    split_amounts = split_large_amounts(extracted_data)

    # Create the invoice table and get the reference to it
    invoice_table = create_invoice_table(doc, split_amounts)

    # Create the job number field below the invoice table
    create_job_number_field(doc, invoice_table)

    # Save the modified document with a new name
    modified_docx_path = docx_path.replace(".docx", "_modified.docx")
    doc.save(modified_docx_path)
    
    return modified_docx_path


def main():
    app = QApplication(sys.argv)
    file_dialog = QFileDialog()
    file_dialog.setNameFilters(["PDF files (*.pdf)"])
    if file_dialog.exec_():
        input_path = file_dialog.selectedFiles()[0]
        with open("input_path.txt", "w") as f:
            f.write(input_path)
        converter = PDFConverter()
        docx_path = converter.convert_pdf_to_docx(input_path)
        if docx_path:
            converter.open_and_edit_docx(docx_path)
            
            # Save and close the initial document before modifications
            save_and_close_initial_doc(converter)

            # Read table data and apply OpenAI extraction
            table_data = read_word_file(docx_path)
            extracted_amounts = extract_data_with_openai(table_data)
            
            # Create a modified document with the new table
            modified_docx_path = create_modified_doc_with_table(docx_path, extracted_amounts)
            
            # Optionally, convert the modified DOCX to PDF
            
            # pdf_path = converter.create_pdf_from_docx(modified_docx_path)
            # if pdf_path:
            #     QMessageBox.information(None, 'Success', f'PDF created: {pdf_path}')
            # else:
            #     QMessageBox.warning(None, 'Error', 'Failed to create PDF from DOCX.')
        else:
            QMessageBox.warning(None, 'Error', 'Failed to convert PDF to DOCX.')
    else:
        QMessageBox.warning(None, 'Error', 'No file selected.')
    return None

if __name__ == "__main__":
    main()

