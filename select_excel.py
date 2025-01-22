import sys
from PyQt5.QtWidgets import QApplication, QFileDialog
import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
import invoice  
import os
from dotenv import load_dotenv
import re

load_dotenv()

logging.basicConfig(level=logging.DEBUG)

def select_excel_file():
    app = QApplication(sys.argv)
    options = QFileDialog.Options()
    options |= QFileDialog.ReadOnly
    file_path, _ = QFileDialog.getOpenFileName(None, "Select an Excel File", "", "Excel Files (*.xlsx);;All Files (*)", options=options)
    if file_path:
        process_selected_excel_file(file_path)
    else:
        logging.error("No Excel file selected")

def process_selected_excel_file(excel_file_path):
    logging.debug(f"Selected Excel file: {excel_file_path}")
    
    # Read the Excel file into a DataFrame
    df = pd.read_excel(excel_file_path)
    
    # Ensure the required columns exist
    required_columns = ["Invoice No.", "TTC Number", "Description", "Amount", "Date"]
    if not all(col in df.columns for col in required_columns):
        logging.error("Excel file is missing one or more required columns")
        return
    
    # Iterate over each row and create a Word document page for each
    extracted_data = []
    for _, row in df.iterrows():
        invoice_number = str(row["Invoice No."])
        ttc_number = str(row["TTC Number"])
        description = str(row["Description"])
        amount = str(row["Amount"])
        date = str(row["Date"])
        
        extracted_data.append((invoice_number, ttc_number, description, amount, date))
    
    # Create the Word document with the extracted data
    create_word_document(extracted_data)


def remove_day_suffix(date_string):
    # This regex finds numbers followed by ST, TH, ND, or RD and removes the suffix
    return re.sub(r'(\d+)(ST|TH|ND|RD)', r'\1', date_string, flags=re.IGNORECASE)

def create_word_document(extracted_data):
    new_doc = docx.Document()

    # Add the rest of the invoice content for each row in extracted_data
    for invoice_number, ttc_number, description, billing, date in extracted_data:
        # Convert description to uppercase and correctly format the date
        description = description.upper()
        date = remove_day_suffix(date).upper()

        # Add and format the header part of the invoice_string
        header_lines = [
            os.getenv("HEADER_LINE_1"),
            os.getenv("HEADER_LINE_2"),
            os.getenv("HEADER_LINE_3"),
            os.getenv("HEADER_LINE_4")
        ]

        for line in header_lines:
            header_paragraph = new_doc.add_paragraph(line)
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment for header
            header_run = header_paragraph.runs[0]
            header_run.font.size = Pt(11)  # Set font size to 11 for the header
            header_run.font.name = 'Courier'  # Set font to Courier

            # Set line spacing to single
            header_paragraph.paragraph_format.line_spacing = 1

        # Add an empty paragraph for spacing
        new_doc.add_paragraph('')

        # Replace placeholders with data
        page_content = invoice.invoice_string.replace('<<invoice>>', invoice_number)
        page_content = page_content.replace('<<job>>', ttc_number)
        page_content = page_content.replace('<<description>>', description)
        page_content = page_content.replace('<<billing>>', billing)
        page_content = page_content.replace('<<date>>', date)

        # Skip the header part that was already added and process the rest of the content
        lines = page_content.split('\n')[5:]
        for line in lines:
            para = new_doc.add_paragraph(line)

            # Align "INVOICE NO.: <<invoice>>" and "DATE: <<date>>" to the right
            if "INVOICE NO." in line or "DATE:" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right alignment for invoice and date
            elif "THANK YOU" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment for THANK YOU
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left alignment for other content

            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)  # Set font size to 9 for the content
                run.font.name = 'Courier'  # Set font to Courier

            # Set line spacing to single
            para.paragraph_format.line_spacing = 1

        # Add a page break after processing each row's content
        new_doc.add_page_break()

    # Ensure the /output directory exists
    output_dir = os.path.join(os.getcwd(), 'final invoice output')
    os.makedirs(output_dir, exist_ok=True)

    # Save the document in the /output directory
    output_path = os.path.join(output_dir, 'Formatted_Invoices_From_Excel.docx')
    new_doc.save(output_path)
    logging.info(f"Formatted document saved as {output_path}")
    
if __name__ == "__main__":
    select_excel_file()

