import sys
import os
import re
import logging
import sqlite3
import datetime  # For generating batch IDs
from dotenv import load_dotenv
from PyQt5.QtWidgets import QApplication, QFileDialog, QInputDialog
from email import policy
from email.parser import BytesParser
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import dspy
import invoice  # Ensure your invoice template module is imported

from pdf_to_docx_ import PDFConverter

load_dotenv()
logging.basicConfig(level=logging.DEBUG)

# Initialize Qt Application for dialogs
app = QApplication(sys.argv)

def get_user_date():
    """
    Prompt the user for the desired date format string.
    """
    text, ok = QInputDialog.getText(None, "Date Format Input", 
                                    "Enter desired date format (e.g., January 15, 2024):")
    if ok and text:
        return text.strip()
    else:
        # Fallback default
        return "January 15, 2024"

# Prompt user for the desired date format before configuring DSPy
desired_date_format = get_user_date()

# Configure DSPy with your OpenAI API key
dspy.configure(lm=dspy.LM('openai/gpt-4o'))


def refine_city_name(market):
    try:
        response = city_extractor(text=market)
        return response.city.strip()
    except Exception as e:
        logging.warning(f"Failed to refine city name: {e}")
        return market




# Define DSPy signature with the user-provided date format
class ExtractInvoiceInfo(dspy.Signature):
    """
    Extract invoice information, including market (city name only).

    """.format(date_format=desired_date_format)
    
    text: str = dspy.InputField()
    invoices: list[dict[str, str]] = dspy.OutputField(
        desc="List of invoices, each with keys: Invoice No., TTC Number, Description, Amount, Date, Market"
    )


# Initialize the DSPy prediction module for extracting invoice info
invoice_extractor = dspy.Predict(ExtractInvoiceInfo)



class ExtractCityOnly(dspy.Signature):
    """
    Extract only the city name from a market field.
    """
    text: str = dspy.InputField()
    city: str = dspy.OutputField(desc="City name extracted from market field")

city_extractor = dspy.Predict(ExtractCityOnly)




def select_eml_file():
    options = QFileDialog.Options()
    options |= QFileDialog.ReadOnly
    file_path, _ = QFileDialog.getOpenFileName(None, "Select an EML File", "", 
                                               "Email Files (*.eml);;All Files (*)", 
                                               options=options)
    if file_path:
        process_selected_eml_file(file_path)
    else:
        logging.error("No EML file selected")

def process_selected_eml_file(eml_file_path):
    logging.debug(f"Selected EML file: {eml_file_path}")

    # Parse the .eml file
    with open(eml_file_path, 'rb') as fp:
        msg = BytesParser(policy=policy.default).parse(fp)

    # Check and save attachments if any
    attachment_dir = os.path.join(os.getcwd(), 'downloaded files email')
    os.makedirs(attachment_dir, exist_ok=True)
    for part in msg.walk():
        content_disposition = part.get("Content-Disposition", "")
        if "attachment" in content_disposition:
            filename = part.get_filename()
            if filename:
                file_data = part.get_payload(decode=True)
                file_path = os.path.join(attachment_dir, filename)
                with open(file_path, 'wb') as f:
                    f.write(file_data)
                logging.info(f"Attachment {filename} saved to {attachment_dir}.")

    # Extract the plain text body of the email
    email_body = None
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == 'text/plain':
                email_body = part.get_content()
                break
    else:
        email_body = msg.get_content()

    if not email_body:
        logging.error("No plain text content found in the email.")
        return

    # Use DSPy to extract structured invoice data
    extracted_data = extract_structured_data_from_email(email_body)

    if not extracted_data:
        logging.error("Failed to extract data from email via DSPy.")
        return

    # Save extracted data into SQLite database with a batch ID
    save_invoices_to_db(extracted_data)

    # Create a Word document using the extracted structured data
    create_word_document(extracted_data)

    convert_pdf_attachments(attachment_dir)



def extract_structured_data_from_email(email_body):
    try:
        # Use DSPy to extract invoice information
        response = invoice_extractor(text=email_body)

        # Extract the list of invoices from the DSPy response
        structured_data = response.invoices

        # Convert list of dicts to list of tuples in expected order
        extracted_data = [
            (
                invoice.get("Invoice No.", ""),
                invoice.get("TTC Number", ""),
                invoice.get("Description", ""),
                invoice.get("Amount", ""),
                invoice.get("Date", ""),
                refine_city_name(invoice.get("Market", "")) # Extract the market (city name) if provided
            )
            for invoice in structured_data
        ]

        # Filter out duplicate invoices based on numeric portion only
        seen_numbers = set()
        unique_data = []
        for invoice_no, ttc_number, description, amount, date, market in extracted_data:
            # Extract numeric portion of the invoice number
            numeric_part = re.match(r'(\d+)', invoice_no.strip())
            if numeric_part:
                num = numeric_part.group(1)
                if num not in seen_numbers:
                    seen_numbers.add(num)
                    unique_data.append((invoice_no, ttc_number, description, amount, date, market))
                else:
                    logging.warning(f"Duplicate invoice found and skipped: {invoice_no}")
            else:
                # If no numeric part, include the invoice as-is
                unique_data.append((invoice_no, ttc_number, description, amount, date, market))

        return unique_data
    except Exception as e:
        logging.error(f"Error during DSPy extraction: {e}")
        return None

def remove_day_suffix(date_string):
    # This regex finds numbers followed by ST, TH, ND, or RD and removes the suffix
    return re.sub(r'(\d+)(ST|TH|ND|RD)', r'\1', date_string, flags=re.IGNORECASE)

def create_word_document(extracted_data):
    new_doc = docx.Document()

    # Unpack only the first five values and ignore the sixth (market)
    for invoice_number, ttc_number, description, amount, date, _ in extracted_data:
        description = description.upper()
        date = remove_day_suffix(date).upper()

        header_lines = [
            os.getenv("HEADER_LINE_1"),
            os.getenv("HEADER_LINE_2"),
            os.getenv("HEADER_LINE_3"),
            os.getenv("HEADER_LINE_4")
        ]

        for line in header_lines:
            header_paragraph = new_doc.add_paragraph(line)
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = header_paragraph.runs[0]
            header_run.font.size = Pt(11)
            header_run.font.name = 'Courier'
            header_paragraph.paragraph_format.line_spacing = 1

        new_doc.add_paragraph('')

        page_content = invoice.invoice_string.replace('<<invoice>>', invoice_number)
        page_content = page_content.replace('<<job>>', ttc_number)
        page_content = page_content.replace('<<description>>', description)
        page_content = page_content.replace('<<billing>>', amount)
        page_content = page_content.replace('<<date>>', date)

        lines = page_content.split('\n')[5:]
        for line in lines:
            para = new_doc.add_paragraph(line)

            if "INVOICE NO." in line or "DATE:" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif "THANK YOU" in line:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = 'Courier'

            para.paragraph_format.line_spacing = 1

        new_doc.add_page_break()

    output_dir = os.path.join(os.getcwd(), 'final invoice output')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'Formatted_Invoices_From_EML.docx')
    new_doc.save(output_path)
    logging.info(f"Formatted document saved as {output_path}")

def save_invoices_to_db(invoices):
    try:
        # Connect to (or create) the SQLite database
        conn = sqlite3.connect('invoices.db')
        cursor = conn.cursor()

        # Create table if it doesn't exist, including the 'market' column
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                batch_id TEXT,
                invoice_no TEXT,
                ttc_number TEXT,
                description TEXT,
                amount TEXT,
                date TEXT,
                market TEXT
            );
        """)

        # Create a unique batch_id for this run (e.g., timestamp)
        batch_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        # Insert each invoice record with the batch_id into the table
        for invoice_no, ttc_number, description, amount, date, market in invoices:
            cursor.execute("""
                INSERT INTO invoices (batch_id, invoice_no, ttc_number, description, amount, date, market)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (batch_id, invoice_no, ttc_number, description, amount, date, market))

        # Commit changes and close connection
        conn.commit()
        conn.close()
        logging.info(f"Invoice data saved to SQLite database with batch_id {batch_id}.")
    except Exception as e:
        logging.error(f"Failed to save invoices to database: {e}")

def convert_pdf_attachments(attachment_dir):
    """
    Iterate over PDF files in the given directory and convert them to DOCX using PDFConverter.
    """
    pdf_converter = PDFConverter()
    for filename in os.listdir(attachment_dir):
        if filename.lower().endswith('.pdf'):
            pdf_path = os.path.join(attachment_dir, filename)
            logging.info(f"Converting {pdf_path} to DOCX...")
            output_file = pdf_converter.convert_pdf_to_docx(pdf_path)
            if output_file:
                logging.info(f"Successfully converted {filename} to {output_file}")
            else:
                logging.error(f"Failed to convert {filename}")



if __name__ == "__main__":
    select_eml_file()
    sys.exit(app.exec_())  # Cleanly exit Qt application
